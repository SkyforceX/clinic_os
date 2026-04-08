
from __future__ import annotations

from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Pt, RGBColor


CONTRACT_CATALOG_TABLE_MARKER = "{{ CONTRACT_CATALOG_TABLE }}"

QUOTATION_TABLE_MARKER = "{{ QUOTATION_TABLE }}"


def _iter_all_tables(parent):
    for table in getattr(parent, "tables", []):
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_tables(cell)


def _iter_all_paragraphs(parent):
    for paragraph in getattr(parent, "paragraphs", []):
        yield paragraph
    for table in getattr(parent, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def _replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    if not full_text:
        return

    replaced = full_text
    for key, value in replacements.items():
        replaced = replaced.replace(key, str(value or ""))

    if replaced == full_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = replaced


def _replace_text_everywhere(doc: Document, replacements: dict):
    for paragraph in _iter_all_paragraphs(doc):
        _replace_text_in_paragraph(paragraph, replacements)

    for section in doc.sections:
        for paragraph in _iter_all_paragraphs(section.header):
            _replace_text_in_paragraph(paragraph, replacements)
        for paragraph in _iter_all_paragraphs(section.footer):
            _replace_text_in_paragraph(paragraph, replacements)


def _remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
    paragraph._p = paragraph._element = None


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def _set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _clear_cell(cell):
    cell._tc.clear_content()
    cell.add_paragraph("")


def _style_runs(paragraph, *, bold=False, italic=False, size=10.5, color=None, font_name="Arial"):
    for run in paragraph.runs:
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), font_name)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _write_cell_text(
    cell,
    text,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    bold=False,
    italic=False,
    size=10.0,
    color=None,
    shading=None,
):
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.add_run(str(text or ""))
    _style_runs(p, bold=bold, italic=italic, size=size, color=color)
    _set_paragraph_spacing(p, before=0, after=0, line=1.12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell, top=60, start=70, bottom=60, end=70)
    _set_cell_border(
        cell,
        top={"val": "single", "sz": 8, "color": "1F1F1F"},
        bottom={"val": "single", "sz": 8, "color": "1F1F1F"},
        left={"val": "single", "sz": 8, "color": "1F1F1F"},
        right={"val": "single", "sz": 8, "color": "1F1F1F"},
    )
    if shading:
        _set_cell_shading(cell, shading)


def _set_row_repeat_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _set_table_width(table, pct: int = 100):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct * 50))


def _set_table_fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _set_cell_width(cell, width_cm: float):
    width_twips = int(width_cm * 567)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Cm(width_cm)


def _set_table_grid(table, widths_cm: list[float]):
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)

    tbl_grid = OxmlElement("w:tblGrid")
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width_cm * 567)))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)


def _set_table_borders(table, size=4, color="888888"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{border_name}"
        element = tbl_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tbl_borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))   # 4 = mỏng đẹp
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _move_table_after_paragraph(paragraph, table):
    paragraph._p.addnext(table._tbl)


def _find_table_marker(doc: Document):
    for paragraph in doc.paragraphs:
        if QUOTATION_TABLE_MARKER in paragraph.text:
            return ("document", paragraph)

    for table in _iter_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if QUOTATION_TABLE_MARKER in paragraph.text:
                        return ("cell", cell, paragraph)

    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in _iter_all_paragraphs(header_footer):
                if QUOTATION_TABLE_MARKER in paragraph.text:
                    return ("header_footer", paragraph)

    return None


def _build_replacements(payload: dict) -> dict:
    quotation = payload["quotation"]
    totals = payload["totals"]
    today = date.today()

    return {
        "{{ issue_day }}": str(today.day),
        "{{ issue_month }}": str(today.month),
        "{{ issue_year }}": str(today.year),
        "{{ issue_date }}": today.strftime("%d/%m/%Y"),
        "{{ quotation_id }}": quotation["id"],
        "{{ quotation_number }}": quotation["id"],
        "{{ company_name }}": quotation["company_name"],
        "{{ contact_name }}": quotation["contact_name"],
        "{{ company_address }}": quotation["company_address"],
        "{{ valid_until }}": quotation["valid_until_display"],
        "{{ pax_from }}": quotation["pax_from"],
        "{{ male_count }}": quotation["male_count"],
        "{{ female_single_count }}": quotation["female_single_count"],
        "{{ female_family_count }}": quotation["female_family_count"],
        "{{ note }}": quotation["note"],
        "{{ total_male }}": totals["total_male_display"],
        "{{ total_female_single }}": totals["total_female_single_display"],
        "{{ total_female_family }}": totals["total_female_family_display"],
        "{{ grand_total }}": totals["grand_total_display"],
    }


def _add_group_row(table, label: str):
    tr = table.add_row()
    tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    merged = tr.cells[0].merge(tr.cells[5])
    _write_cell_text(
        merged,
        label,
        bold=True,
        size=10.5,
        color="FFFFFF",
        shading="2F5AA8",
    )
    return tr


def _add_item_row(table, line: dict):
    tr = table.add_row()
    tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    service_text = line["item_name"]
    if line["description"]:
        service_text = f"{service_text}\n{line['description']}"

    _write_cell_text(tr.cells[0], line["stt"], align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _write_cell_text(tr.cells[1], service_text, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
    _write_cell_text(
        tr.cells[2],
        line["display_unit_price"] or "",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=bool(line["display_unit_price"]),
        size=10,
        shading="D9E5F6" if line["price_type"] == "standard" else "F6C39A",
    )
    _write_cell_text(tr.cells[3], line["male_mark"], align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _write_cell_text(tr.cells[4], line["female_single_mark"], align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _write_cell_text(tr.cells[5], line["female_family_mark"], align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    return tr


def _append_quotation_table(container, payload: dict):
    table = container.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, size=4, color="999999")
    _set_table_width(table, 100)
    _set_table_fixed_layout(table)

    # fit trong khổ A4 portrait sau khi trừ margin
    widths_cm = [0.95, 7.25, 1.85, 2.55, 2.85, 2.85]
    _set_table_grid(table, widths_cm)

    def _apply_widths(row):
        for idx, width_cm in enumerate(widths_cm):
            _set_cell_width(row.cells[idx], width_cm)

    for row in table.rows:
        _apply_widths(row)

    header = table.rows[0]
    _apply_widths(header)
    _set_row_repeat_as_header(header)
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    headers = [
        "STT",
        "KHÁM TỔNG QUÁT /\nGENERAL EXAMINATION",
        "GIÁ ƯU ĐÃI\n(VND)",
        "NAM",
        "NỮ\nĐỘC THÂN",
        "NỮ\nGIA ĐÌNH",
    ]
    for idx, label in enumerate(headers):
        _write_cell_text(
            header.cells[idx],
            label,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            size=9.5,
            color="FFFFFF",
            shading="2F5AA8",
        )

    for row in payload["display_rows"]:
        if row["row_type"] == "group":
            tr = _add_group_row(table, row["label"])
            _apply_widths(tr)
            continue
        if row["row_type"] == "subgroup":
            tr = _add_group_row(table, row["label"])
            _apply_widths(tr)
            continue
        tr = _add_item_row(table, row["line"])
        _apply_widths(tr)

    totals = payload["totals"]

    # Dòng giá gốc / niêm yết
    tr = table.add_row()
    _apply_widths(tr)
    tr.height = Cm(1.15)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    merged = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged, widths_cm[0] + widths_cm[1] + widths_cm[2])

    _write_cell_text(
        merged,
        "GIÁ TRỊ GÓI KHÁM THEO GIÁ NIÊM YẾT TẠI PHÒNG KHÁM\n(VND/NGƯỜI)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=9.5,
        color="FFFFFF",
        shading="2F5AA8",
    )
    _write_cell_text(
        tr.cells[3],
        totals.get("base_total_male_display") or "0",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        color="FFFFFF",
        shading="2F5AA8",
    )
    _write_cell_text(
        tr.cells[4],
        totals.get("base_total_female_single_display") or "0",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        color="FFFFFF",
        shading="2F5AA8",
    )
    _write_cell_text(
        tr.cells[5],
        totals.get("base_total_female_family_display") or "0",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        color="FFFFFF",
        shading="2F5AA8",
    )
    _set_cell_vertical_center(merged)

    # Dòng giá ưu đãi
    tr = table.add_row()
    _apply_widths(tr)
    tr.height = Cm(0.95)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    merged_discount = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_discount, widths_cm[0] + widths_cm[1] + widths_cm[2])

    _write_cell_text(
        merged_discount,
        "GIÁ ƯU ĐÃI DÀNH CHO QUÝ CTY (VND)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        shading="FFF200",
    )
    _write_cell_text(
        tr.cells[3],
        totals.get("discounted_total_male_display") or "0",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        shading="FFF200",
    )
    _write_cell_text(
        tr.cells[4],
        totals.get("discounted_total_female_single_display") or "0",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        shading="FFF200",
    )
    _write_cell_text(
        tr.cells[5],
        totals.get("discounted_total_female_family_display") or "0",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        shading="FFF200",
    )
    _set_cell_vertical_center(merged_discount)

    # ===== KHÔNG HIỂN THỊ RA PDF "tổng giá trị dự kiến" =====
    # tr = table.add_row()
    # _apply_widths(tr)
    #
    # merged_total = tr.cells[0].merge(tr.cells[2])
    # _set_cell_width(merged_total, widths_cm[0] + widths_cm[1] + widths_cm[2])
    #
    # _write_cell_text(
    #     merged_total,
    #     "TỔNG GIÁ TRỊ HỢP ĐỒNG DỰ KIẾN (VND)",
    #     align=WD_ALIGN_PARAGRAPH.CENTER,
    #     bold=True,
    #     size=10.5,
    #     color="FFFFFF",
    #     shading="2F5AA8",
    # )

    # merged_value = tr.cells[3].merge(tr.cells[5])
    # _set_cell_width(
    #     merged_value,
    #     widths_cm[3] + widths_cm[4] + widths_cm[5],
    # )
    # _write_cell_text(
    #     merged_value,
    #     totals.get("grand_total_display") or "0",
    #     align=WD_ALIGN_PARAGRAPH.CENTER,
    #     bold=True,
    #     size=11,
    #     color="FFFFFF",
    #     shading="2F5AA8",
    # )
    # _set_cell_vertical_center(merged_total)
    # _set_cell_vertical_center(merged_value)

    return table


def _set_cell_vertical_center(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)

    v_align.set(qn("w:val"), "center")


def _inject_table_at_marker(doc: Document, payload: dict) -> bool:
    marker = _find_table_marker(doc)
    if not marker:
        return False

    kind = marker[0]
    if kind == "document":
        paragraph = marker[1]
        paragraph.text = paragraph.text.replace(QUOTATION_TABLE_MARKER, "").strip()
        table = doc.add_table(rows=1, cols=6)
        table._tbl.getparent().remove(table._tbl)
        paragraph._p.addnext(table._tbl)
        # build rows into moved table using lightweight proxy
        proxy = type("TableContainer", (), {"add_table": lambda self, rows, cols: table})()
        _append_quotation_table(proxy, payload)
        if not paragraph.text.strip():
            _remove_paragraph(paragraph)
        return True

    if kind == "cell":
        cell, paragraph = marker[1], marker[2]
        paragraph.text = paragraph.text.replace(QUOTATION_TABLE_MARKER, "").strip()
        _clear_cell(cell)
        _append_quotation_table(cell, payload)
        return True

    return False


def create_default_quotation_docx(payload: dict, output_path: str):
    quotation = payload["quotation"]
    totals = payload["totals"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("BẢNG BÁO GIÁ KHÁM SỨC KHỎE")
    _style_runs(p, bold=True, size=20, color="2F5AA8")
    _set_paragraph_spacing(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PROPOSAL - Medical Check-up")
    _style_runs(p, bold=True, italic=True, size=16, color="F28C38")
    _set_paragraph_spacing(p, after=8)

    for line in [
        f"Kính gửi: {quotation['contact_name']}",
        f"Công ty: {quotation['company_name']}",
        f"Địa chỉ: {quotation['company_address']}",
        f"Giá ưu đãi áp dụng cho số lượng khám dự kiến từ: {quotation['pax_from'] or '___'} người",
        f"Hiệu lực đến: {quotation['valid_until_display'] or ''}",
    ]:
        p = doc.add_paragraph(line)
        _style_runs(p, size=11)
        _set_paragraph_spacing(p, after=2)

    p = doc.add_paragraph()
    p.add_run("Ghi chú: ").bold = True
    p.add_run(quotation["note"] or "")
    _style_runs(p, size=10.5)
    _set_paragraph_spacing(p, after=8)

    _append_quotation_table(doc, payload)

    p = doc.add_paragraph()
    p.add_run("Tổng Nam: ").bold = True
    p.add_run(totals["total_male_display"] or "0")
    _style_runs(p, size=10.5)
    _set_paragraph_spacing(p, before=8, after=2)

    p = doc.add_paragraph()
    p.add_run("Tổng Nữ độc thân: ").bold = True
    p.add_run(totals["total_female_single_display"] or "0")
    _style_runs(p, size=10.5)
    _set_paragraph_spacing(p, after=2)

    p = doc.add_paragraph()
    p.add_run("Tổng Nữ gia đình: ").bold = True
    p.add_run(totals["total_female_family_display"] or "0")
    _style_runs(p, size=10.5)
    _set_paragraph_spacing(p, after=2)

    p = doc.add_paragraph()
    p.add_run("Tổng giá trị dự kiến: ").bold = True
    p.add_run(totals["grand_total_display"] or "0")
    _style_runs(p, size=11)
    _set_paragraph_spacing(p, after=4)

    doc.save(output_path)
    return output_path


def render_quotation_docx(*, payload: dict, output_path: str, template_path: str | None = None):
    if template_path:
        template_file = Path(template_path)
        if template_file.exists() and template_file.is_file():
            doc = Document(str(template_file))
            replacements = _build_replacements(payload)
            _replace_text_everywhere(doc, replacements)
            inserted = _inject_table_at_marker(doc, payload)
            if not inserted:
                doc.add_paragraph("")
                _append_quotation_table(doc, payload)
            doc.save(output_path)
            return output_path

    return create_default_quotation_docx(payload, output_path)


"""
  1. Mở template .docx (upload qua Admin vào DocumentTemplate)
  2. Thay thế tất cả {{ key }} bằng payload value
     → Dùng lại _replace_text_everywhere() đã có
  3. Chèn bảng danh mục vào marker {{ CONTRACT_CATALOG_TABLE }}
     → Bảng giống quotation nhưng không có cột "GIÁ ƯU ĐÃI" tổng hợp cuối
  4. Nếu không có template → tạo fallback text-only
"""

def _build_contract_replacements(payload: dict) -> dict:
    """
    Map tất cả {{ key }} trong .docx template → giá trị thực từ payload.
    Key phải khớp chính xác với text trong file template.
    """
    return {
        # Header
        "{{ contract_number }}":       payload["contract_number"],
        "{{ contract_number_full }}":  payload["contract_number_full"],
        "{{ issue_day }}":             payload["issue_day"],
        "{{ issue_month }}":           payload["issue_month"],
        "{{ issue_year }}":            payload["issue_year"],
        "{{ issue_date }}":            payload["issue_date"],
        "{{ issue_date_vi }}":         payload["issue_date_vi"],

        # Bên B
        "{{ company_name }}":          payload["company_name"],
        "{{ company_address }}":       payload["company_address"],
        "{{ company_phone }}":         payload["company_phone"],
        "{{ company_tax_code }}":      payload["company_tax_code"],
        "{{ signer_b_name }}":         payload["signer_b_name"],
        "{{ signer_b_title }}":        payload["signer_b_title"],

        # Bên A
        "{{ signer_a_name }}":         payload["signer_a_name"],
        "{{ signer_a_title }}":        payload["signer_a_title"],

        # Điều I
        "{{ period_text }}":           payload["period_text"],
        "{{ start_date }}":            payload["start_date"],
        "{{ end_date }}":              payload["end_date"],
        "{{ reception_from_date }}":   payload["reception_from_date"],
        "{{ blood_time_from }}":       payload["blood_time_from"],
        "{{ blood_time_to }}":         payload["blood_time_to"],
        "{{ blood_time_text }}":       payload["blood_time_text"],
        "{{ blood_collection_location }}": payload["blood_collection_location"],

        # Nhân sự
        "{{ male_count }}":            str(payload["male_count"]),
        "{{ female_single_count }}":   str(payload["female_single_count"]),
        "{{ female_family_count }}":   str(payload["female_family_count"]),
        "{{ total_pax }}":             str(payload["total_pax"]),

        # Điều II — Tài chính
        "{{ subtotal_male }}":         payload["subtotal_male"],
        "{{ subtotal_female_single }}": payload["subtotal_female_single"],
        "{{ subtotal_female_family }}": payload["subtotal_female_family"],
        "{{ grand_total }}":           payload["grand_total"],
        "{{ deposit_pct }}":           payload["deposit_pct"],
        "{{ deposit_amount }}":        payload["deposit_amount"],
        "{{ deposit_amount_words }}": payload["deposit_amount_words"],
        "{{ deposit_deadline }}":      payload["deposit_deadline"],
        "{{ settlement_days }}":       payload["settlement_days"],

        # Ghi chú
        "{{ contract_note }}":         payload["contract_note"],
        "{{ note }}":                  payload["note"],
    }


# ── Bảng danh mục dịch vụ (tái sử dụng helper từ quotation) ─────────────────

def _add_contract_catalog_table(container, payload: dict):
    """
    Bảng danh mục dịch vụ cho hợp đồng.
    """
    table = container.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, size=4, color="999999")
    _set_table_width(table, 100)
    _set_table_fixed_layout(table)

    widths_cm = [0.95, 7.25, 1.85, 2.55, 2.85, 2.85]
    _set_table_grid(table, widths_cm)

    def _apply_widths(row):
        for idx, width_cm in enumerate(widths_cm):
            _set_cell_width(row.cells[idx], width_cm)

    for row in table.rows:
        _apply_widths(row)

    header = table.rows[0]
    _apply_widths(header)
    _set_row_repeat_as_header(header)
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    headers = [
        "STT",
        "KHÁM TỔNG QUÁT /\nGENERAL EXAMINATION",
        "GIÁ ƯU ĐÃI\n(VND)",
        "NAM",
        "NỮ\nĐỘC THÂN",
        "NỮ\nGIA ĐÌNH",
    ]
    for idx, label in enumerate(headers):
        _write_cell_text(
            header.cells[idx],
            label,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            size=9.5,
            color="FFFFFF",
            shading="2F5AA8",
        )

    display_rows = payload.get("display_rows") or []

    # Fallback an toàn nếu payload cũ chưa có display_rows
    if not display_rows:
        for line in payload.get("lines", []):
            display_rows.append(
                {
                    "row_type": "item",
                    "group_name": line.get("group_name") or "Khác",
                    "subgroup_name": line.get("subgroup_name") or "",
                    "line": line,
                }
            )

    current_group = None
    current_subgroup = None

    def _add_group_row_local(label: str, *, shading="2F5AA8", italic=False):
        tr = table.add_row()
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        _apply_widths(tr)
        merged = tr.cells[0].merge(tr.cells[5])
        _write_cell_text(
            merged,
            label,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            bold=True,
            italic=italic,
            size=10.5 if not italic else 10,
            color="FFFFFF" if not italic else None,
            shading=shading,
        )
        return tr

    def _mark_text(enabled: bool, price_type: str, note: str | None = None) -> str:
        if not enabled:
            return "–"
        if price_type == "free":
            return "Miễn phí"
        if price_type == "gift":
            return note or "TẶNG"
        return "✓"

    def _add_item_row_local(line: dict):
        tr = table.add_row()
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        _apply_widths(tr)

        service_text = line.get("item_name") or ""
        description = line.get("description") or ""
        if description:
            service_text = f"{service_text}\n{description}"

        price_type = line.get("price_type") or "standard"
        note = line.get("note") or ""
        price_label = line.get("display_unit_price") or ""

        price_shading = "D9E5F6" if price_type == "standard" else "F6C39A"
        mark_shading = "FDEBD0" if price_type in ("free", "gift") else None

        _write_cell_text(
            tr.cells[0],
            line.get("stt") or "",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=10,
        )
        _write_cell_text(
            tr.cells[1],
            service_text,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            size=10,
        )
        _write_cell_text(
            tr.cells[2],
            price_label,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=bool(price_label),
            size=10,
            shading=price_shading,
        )
        _write_cell_text(
            tr.cells[3],
            line.get("male_mark")
            or _mark_text(bool(line.get("for_male")), price_type, note),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=10,
            shading=mark_shading,
        )
        _write_cell_text(
            tr.cells[4],
            line.get("female_single_mark")
            or _mark_text(bool(line.get("for_female_single")), price_type, note),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=10,
            shading=mark_shading,
        )
        _write_cell_text(
            tr.cells[5],
            line.get("female_family_mark")
            or _mark_text(bool(line.get("for_female_family")), price_type, note),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=10,
            shading=mark_shading,
        )
        return tr

    for row in display_rows:
        row_type = row.get("row_type")

        if row_type == "group":
            current_group = row.get("label") or "Khác"
            current_subgroup = None
            _add_group_row_local(current_group, shading="2F5AA8", italic=False)
            continue

        if row_type == "subgroup":
            current_subgroup = row.get("label") or ""
            if current_subgroup:
                _add_group_row_local(current_subgroup, shading="D6EAF8", italic=True)
            continue

        line = row.get("line") or {}
        group_name = (line.get("group_name") or row.get("group_name") or "Khác").strip()
        subgroup_name = (line.get("subgroup_name") or row.get("subgroup_name") or "").strip()

        if group_name and group_name != current_group:
            current_group = group_name
            current_subgroup = None
            _add_group_row_local(current_group, shading="2F5AA8", italic=False)

        if subgroup_name and subgroup_name != current_subgroup:
            current_subgroup = subgroup_name
            _add_group_row_local(current_subgroup, shading="D6EAF8", italic=True)

        _add_item_row_local(line)
    
    # Row giá niêm yết / người
    tr = table.add_row()
    _apply_widths(tr)
    tr.height = Cm(0.95)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    merged = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged, widths_cm[0] + widths_cm[1] + widths_cm[2])
    
    _write_cell_text(
        merged,
        "GIÁ TRỊ GÓI KHÁM THEO GIÁ NIÊM YẾT (VND/NGƯỜI)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        italic=False,
        size=9.5,
        color="D35400",
        shading="F3F3F3",
    )
    _write_cell_text(
        tr.cells[3],
        payload.get("base_total_male_display") or "–",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        italic=True,
        size=10,
        color="D35400",
        shading="F3F3F3",
    )
    _write_cell_text(
        tr.cells[4],
        payload.get("base_total_female_single_display") or "–",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        italic=True,
        size=10,
        color="D35400",
        shading="F3F3F3",
    )
    _write_cell_text(
        tr.cells[5],
        payload.get("base_total_female_family_display") or "–",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        italic=True,
        size=10,
        color="D35400",
        shading="F3F3F3",
    )
    
    # Row giá ưu đãi
    tr = table.add_row()
    _apply_widths(tr)
    tr.height = Cm(0.95)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    merged = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged, widths_cm[0] + widths_cm[1] + widths_cm[2])
    
    _write_cell_text(
        merged,
        "GIÁ ƯU ĐÃI DÀNH CHO QUÝ CÔNG TY",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        color="FFFFFF",
        shading="2F7F8E",
    )
    _write_cell_text(
        tr.cells[3],
        payload.get("subtotal_male") or "–",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        color="FFFFFF",
        shading="2F7F8E",
    )
    _write_cell_text(
        tr.cells[4],
        payload.get("subtotal_female_single") or "–",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        color="FFFFFF",
        shading="2F7F8E",
    )
    _write_cell_text(
        tr.cells[5],
        payload.get("subtotal_female_family") or "–",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        color="FFFFFF",
        shading="2F7F8E",
    )
    
    # Row số lượng nhân viên
    tr = table.add_row()
    _apply_widths(tr)
    tr.height = Cm(0.95)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    merged = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged, widths_cm[0] + widths_cm[1] + widths_cm[2])
    
    _write_cell_text(
        merged,
        "SỐ LƯỢNG NHÂN VIÊN",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        shading="FAFAFA",
    )
    _write_cell_text(
        tr.cells[3],
        str(payload.get("male_count") or 0),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        shading="FAFAFA",
    )
    _write_cell_text(
        tr.cells[4],
        str(payload.get("female_single_count") or 0),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        shading="FAFAFA",
    )
    _write_cell_text(
        tr.cells[5],
        str(payload.get("female_family_count") or 0),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10,
        shading="FAFAFA",
    )
    
    # Row giá trị hợp đồng
    tr = table.add_row()
    _apply_widths(tr)
    
    merged_label = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_label, widths_cm[0] + widths_cm[1] + widths_cm[2])
    
    _write_cell_text(
        merged_label,
        "GIÁ TRỊ HỢP ĐỒNG (VND)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=10.5,
        shading="F3F3F3",
    )
    
    merged_value = tr.cells[3].merge(tr.cells[5])
    _set_cell_width(
        merged_value,
        widths_cm[3] + widths_cm[4] + widths_cm[5],
    )
    _write_cell_text(
        merged_value,
        payload.get("grand_total") or "0",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=11,
        shading="F3F3F3",
    )
    _set_cell_vertical_center(merged_label)
    _set_cell_vertical_center(merged_value)
    
    return table


def _find_contract_catalog_marker(doc: Document):
    """Tìm đoạn văn chứa {{ CONTRACT_CATALOG_TABLE }} trong toàn bộ document."""
    for paragraph in doc.paragraphs:
        if CONTRACT_CATALOG_TABLE_MARKER in paragraph.text:
            return ("document", paragraph)

    for table in _iter_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if CONTRACT_CATALOG_TABLE_MARKER in paragraph.text:
                        return ("cell", cell, paragraph)

    return None


def _inject_contract_catalog(doc: Document, payload: dict) -> bool:
    """Inject bảng danh mục tại vị trí marker."""
    marker = _find_contract_catalog_marker(doc)
    if not marker:
        return False

    kind = marker[0]

    if kind == "document":
        paragraph = marker[1]
        paragraph.text = paragraph.text.replace(CONTRACT_CATALOG_TABLE_MARKER, "").strip()

        # Tạo table rỗng rồi di chuyển sau paragraph
        table = doc.add_table(rows=1, cols=6)   # sửa 5 -> 6
        table._tbl.getparent().remove(table._tbl)
        paragraph._p.addnext(table._tbl)

        proxy = type("TC", (), {"add_table": lambda self, rows, cols: table})()
        _add_contract_catalog_table(proxy, payload)

        if not paragraph.text.strip():
            _remove_paragraph(paragraph)
        return True

    if kind == "cell":
        cell, paragraph = marker[1], marker[2]
        paragraph.text = paragraph.text.replace(CONTRACT_CATALOG_TABLE_MARKER, "").strip()
        _clear_cell(cell)
        _add_contract_catalog_table(cell, payload)
        return True

    return False


# ── Fallback: tạo docx đơn giản khi không có template ───────────────────────

def create_default_contract_docx(payload: dict, output_path: str) -> str:
    """Tạo docx tối giản khi không có template upload."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HỢP ĐỒNG KHÁM SỨC KHỎE ĐỊNH KỲ")
    _style_runs(p, bold=True, size=16, color="1A5276")
    _set_paragraph_spacing(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Số: {payload['contract_number_full']}")
    _style_runs(p, size=12)
    _set_paragraph_spacing(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Lập {payload['issue_date_vi']}")
    _style_runs(p, italic=True, size=11)
    _set_paragraph_spacing(p, after=12)

    # Bên B
    for line in [
        f"BÊN B: {payload['company_name']}",
        f"Địa chỉ: {payload['company_address']}",
        f"Điện thoại: {payload['company_phone']}",
        f"Đại diện: {payload['signer_b_name']} – {payload['signer_b_title']}",
        f"Mã số thuế: {payload['company_tax_code']}",
    ]:
        if line.strip().endswith(":") or line.strip().endswith(": "):
            continue
        p = doc.add_paragraph(line)
        _style_runs(p, size=11)
        _set_paragraph_spacing(p, after=2)

    doc.add_paragraph()

    # Điều I
    p = doc.add_paragraph("ĐIỀU I: THỜI GIAN VÀ ĐỊA ĐIỂM THỰC HIỆN HỢP ĐỒNG")
    _style_runs(p, bold=True, size=12)
    _set_paragraph_spacing(p, before=8, after=4)

    p = doc.add_paragraph(f"Thời gian thực hiện: {payload['period_text']}")
    _style_runs(p, size=11)
    if payload["blood_time_text"]:
        p = doc.add_paragraph(f"Giờ lấy mẫu: {payload['blood_time_text']}")
        _style_runs(p, size=11)
    if payload["blood_collection_location"]:
        p = doc.add_paragraph(f"Địa điểm lấy mẫu: {payload['blood_collection_location']}")
        _style_runs(p, size=11)

    # Bảng danh mục
    doc.add_paragraph()
    p = doc.add_paragraph("DANH MỤC KHÁM")
    _style_runs(p, bold=True, size=11)
    _set_paragraph_spacing(p, after=4)

    if payload["lines"]:
        _add_contract_catalog_table(doc, payload)

    doc.add_paragraph()

    # Điều II
    p = doc.add_paragraph("ĐIỀU II: GIÁ TRỊ HỢP ĐỒNG VÀ PHƯƠNG THỨC THANH TOÁN")
    _style_runs(p, bold=True, size=12)
    _set_paragraph_spacing(p, before=8, after=4)

    for line in [
        f"Tổng giá trị hợp đồng: {payload['grand_total']} VNĐ",
        f"Đặt cọc {payload['deposit_pct']}: {payload['deposit_amount']} VNĐ",
        f"Hạn đặt cọc: {payload['deposit_deadline']}",
        f"Quyết toán trong vòng {payload['settlement_days']} ngày sau khi bàn giao kết quả.",
    ]:
        p = doc.add_paragraph(line)
        _style_runs(p, size=11)
        _set_paragraph_spacing(p, after=3)

    if payload["contract_note"]:
        doc.add_paragraph()
        p = doc.add_paragraph(f"Ghi chú: {payload['contract_note']}")
        _style_runs(p, size=11)

    # Ký kết
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("BÊN A" + " " * 40 + "BÊN B")
    _style_runs(p, bold=True, size=12)

    doc.save(output_path)
    return output_path


# ── Entry point chính ────────────────────────────────────────────────────────

def render_contract_docx(
    *,
    payload: dict,
    output_path: str,
    template_path: str | None = None,
) -> str:
    """
    Render hợp đồng ra file .docx.

    - Nếu có template_path → mở template, fill {{ key }}, chèn bảng tại marker
    - Nếu không             → tạo fallback docx đơn giản
    """
    if template_path:
        template_file = Path(template_path)
        if template_file.exists() and template_file.is_file():
            doc = Document(str(template_file))
            replacements = _build_contract_replacements(payload)
            _replace_text_everywhere(doc, replacements)

            # Chèn bảng danh mục nếu có marker
            if payload.get("lines"):
                _inject_contract_catalog(doc, payload)

            doc.save(output_path)
            return output_path

    return create_default_contract_docx(payload, output_path)