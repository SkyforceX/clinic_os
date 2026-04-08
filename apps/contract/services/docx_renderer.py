from __future__ import annotations

import re
from copy import deepcopy
from datetime import date
from pathlib import Path

from html import unescape
from html.parser import HTMLParser

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CONTRACT_CATALOG_TABLE_MARKER = "{{ CONTRACT_CATALOG_TABLE }}"
QUOTATION_TABLE_MARKER = "{{ QUOTATION_TABLE }}"


def _to_int(value) -> int:
    """Local helper — mirror of document_payloads._to_int để tránh circular import."""
    if value in (None, "", False):
        return 0
    try:
        return int(value)
    except Exception:
        try:
            return int(str(value).replace(",", "").replace(".", ""))
        except Exception:
            return 0


def _iter_all_tables(parent):
    for table in getattr(parent, "tables", []):
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_tables(cell)


def _iter_all_paragraphs(parent):
    for paragraph in getattr(parent, "paragraphs", []):
        yield paragraph
    for table in getattr(parent, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def _replace_text_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    if not full_text:
        return
    replaced = full_text
    for key, value in replacements.items():
        replaced = replaced.replace(key, str(value or ""))
    if replaced == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = replaced


def _replace_text_everywhere(doc: Document, replacements: dict):
    for paragraph in _iter_all_paragraphs(doc):
        _replace_text_in_paragraph(paragraph, replacements)
    for section in doc.sections:
        for paragraph in _iter_all_paragraphs(section.header):
            _replace_text_in_paragraph(paragraph, replacements)
        for paragraph in _iter_all_paragraphs(section.footer):
            _replace_text_in_paragraph(paragraph, replacements)


def _remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def _set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _clear_cell(cell):
    cell._tc.clear_content()
    cell.add_paragraph("")


def _style_runs(paragraph, *, bold=False, italic=False, size=10.5, color=None, font_name="Arial"):
    for run in paragraph.runs:
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), font_name)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _write_cell_text(cell, text, *, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False,
                     size=10.0, color=None, shading=None):
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.add_run(str(text or ""))
    _style_runs(p, bold=bold, italic=italic, size=size, color=color)
    _set_paragraph_spacing(p, before=0, after=0, line=1.12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell, top=60, start=70, bottom=60, end=70)
    _set_cell_border(
        cell,
        top={"val": "single", "sz": 8, "color": "1F1F1F"},
        bottom={"val": "single", "sz": 8, "color": "1F1F1F"},
        left={"val": "single", "sz": 8, "color": "1F1F1F"},
        right={"val": "single", "sz": 8, "color": "1F1F1F"},
    )
    if shading:
        _set_cell_shading(cell, shading)


def _set_row_repeat_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _set_table_width(table, pct: int = 100):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct * 50))


def _set_table_fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _set_cell_width(cell, width_cm: float):
    width_twips = int(width_cm * 567)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Cm(width_cm)


def _set_table_grid(table, widths_cm: list[float]):
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width_cm * 567)))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)


def _set_table_borders(table, size=4, color="888888"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{border_name}"
        element = tbl_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_vertical_center(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        tc_pr.append(v_align)
    v_align.set(qn("w:val"), "center")


def _find_table_marker(doc: Document):
    for paragraph in doc.paragraphs:
        if QUOTATION_TABLE_MARKER in paragraph.text:
            return ("document", paragraph)
    for table in _iter_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if QUOTATION_TABLE_MARKER in paragraph.text:
                        return ("cell", cell, paragraph)
    return None


def _build_replacements(payload: dict) -> dict:
    quotation = payload["quotation"]
    totals = payload.get("totals", {})
    today = date.today()
    return {
        "{{ issue_day }}": str(today.day),
        "{{ issue_month }}": str(today.month),
        "{{ issue_year }}": str(today.year),
        "{{ issue_date }}": today.strftime("%d/%m/%Y"),
        "{{ quotation_id }}": quotation.get("id", ""),
        "{{ quotation_number }}": quotation.get("id", ""),
        "{{ company_name }}": quotation.get("company_name", ""),
        "{{ contact_name }}": quotation.get("contact_name", ""),
        "{{ company_address }}": quotation.get("company_address", ""),
        "{{ valid_until }}": quotation.get("valid_until_display", ""),
        "{{ pax_from }}": quotation.get("pax_from", ""),
        "{{ note }}": quotation.get("note", ""),
        "{{ total_male }}": totals.get("total_male_display", ""),
        "{{ total_female_single }}": totals.get("total_female_single_display", ""),
        "{{ total_female_family }}": totals.get("total_female_family_display", ""),
        "{{ grand_total }}": totals.get("grand_total_display", ""),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Multi-package table builder
# ─────────────────────────────────────────────────────────────────────────────

def _compute_col_widths(num_cols: int) -> list[float]:
    """
    Tính widths_cm cho bảng N cột đối tượng.
    Cột STT: 0.95, Cột dịch vụ: 7.25 (fixed), cột giá niêm yết: 1.85 (fixed),
    Còn lại chia đều cho num_cols cột đối tượng.
    """
    stt_w = 0.95
    svc_w = 7.25
    price_w = 1.85
    available = 18.3 - stt_w - svc_w - price_w  # A4 portrait - margins ~18.3cm
    col_w = round(available / max(num_cols, 1), 2)
    return [stt_w, svc_w, price_w] + [col_w] * num_cols


def _append_package_table(insert_after_element, doc: Document, pkg_payload: dict):
    """
    Chèn bảng giá của một gói khám vào document SAU insert_after_element.
    Trả về phần tử XML cuối cùng được chèn (để chain tiếp).
    """
    columns = pkg_payload.get("columns_json") or [
        {"key": "male", "label": "NAM"},
        {"key": "female_single", "label": "NỮ ĐỘC THÂN"},
        {"key": "female_family", "label": "NỮ GIA ĐÌNH"},
    ]
    num_cols = len(columns)
    total_cols = 3 + num_cols  # STT + Dịch vụ + Giá NL + N cột

    widths_cm = _compute_col_widths(num_cols)

    # Tiêu đề gói khám (paragraph)
    pkg_para = OxmlElement("w:p")
    pkg_pPr = OxmlElement("w:pPr")
    pkg_spacing = OxmlElement("w:spacing")
    pkg_spacing.set(qn("w:before"), "120")
    pkg_spacing.set(qn("w:after"), "60")
    pkg_pPr.append(pkg_spacing)
    pkg_para.append(pkg_pPr)
    pkg_r = OxmlElement("w:r")
    pkg_rPr = OxmlElement("w:rPr")
    pkg_b = OxmlElement("w:b")
    pkg_sz = OxmlElement("w:sz")
    pkg_sz.set(qn("w:val"), "22")  # 11pt
    pkg_color = OxmlElement("w:color")
    pkg_color.set(qn("w:val"), "1A5276")
    pkg_rPr.append(pkg_b)
    pkg_rPr.append(pkg_sz)
    pkg_rPr.append(pkg_color)
    pkg_r.append(pkg_rPr)
    pkg_t = OxmlElement("w:t")
    pkg_t.text = pkg_payload.get("name", "Gói khám")
    pkg_r.append(pkg_t)
    pkg_para.append(pkg_r)
    insert_after_element.addnext(pkg_para)

    # Tạo bảng
    table = doc.add_table(rows=1, cols=total_cols)
    table._tbl.getparent().remove(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, size=4, color="999999")
    _set_table_width(table, 100)
    _set_table_fixed_layout(table)
    _set_table_grid(table, widths_cm)

    def apply_widths(row):
        for idx, width_cm in enumerate(widths_cm):
            _set_cell_width(row.cells[idx], width_cm)

    for row in table.rows:
        apply_widths(row)

    # Header row
    header = table.rows[0]
    apply_widths(header)
    _set_row_repeat_as_header(header)
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    header_labels = ["STT", "KHÁM TỔNG QUÁT /\nGENERAL EXAMINATION", "GIÁ NIÊM YẾT\n(VND)"]
    for col in columns:
        header_labels.append(col.get("label", "").upper())

    for idx, label in enumerate(header_labels):
        _write_cell_text(
            header.cells[idx], label,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True, size=9.5, color="FFFFFF", shading="2F5AA8",
        )

    # Data rows
    for row_data in (pkg_payload.get("display_rows") or []):
        row_type = row_data.get("row_type")
        if row_type in ("group", "subgroup"):
            tr = table.add_row()
            apply_widths(tr)
            tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            merged = tr.cells[0].merge(tr.cells[total_cols - 1])
            is_sub = row_type == "subgroup"
            _write_cell_text(
                merged, row_data.get("label", ""),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                bold=True, italic=is_sub,
                size=10 if is_sub else 10.5,
                color=None if is_sub else "FFFFFF",
                shading="D6EAF8" if is_sub else "2F5AA8",
            )
            continue

        line = row_data.get("line") or {}
        tr = table.add_row()
        apply_widths(tr)
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        price_type = line.get("price_type") or "standard"
        note = line.get("note") or ""

        # STT
        _write_cell_text(tr.cells[0], line.get("stt", ""), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

        # Dịch vụ
        svc_text = line.get("item_name", "")
        if line.get("description"):
            svc_text = f"{svc_text}\n{line['description']}"
        _write_cell_text(tr.cells[1], svc_text, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)

        # Giá niêm yết
        list_price = line.get("list_price") or 0
        if price_type == "free":
            price_display = "Miễn phí"
        elif price_type == "gift":
            price_display = "TẶNG"
        else:
            from apps.contract.services.document_payloads import fmt_vnd
            price_display = fmt_vnd(list_price) if list_price else "—"
        _write_cell_text(tr.cells[2], price_display, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bool(list_price), size=10, shading="D9E5F6")

        # Cột đối tượng
        from apps.contract.services.document_payloads import fmt_vnd
        for ci, col in enumerate(columns):
            key = col.get("key", "")
            cell_idx = 3 + ci
            if price_type in ("free", "gift"):
                cell_text = note or ("Miễn phí" if price_type == "free" else "TẶNG")
            elif key == "male":
                val = line.get("price_male") or 0
                cell_text = fmt_vnd(val) if val else "—"
            elif key == "female_single":
                val = line.get("price_female_single") or 0
                cell_text = fmt_vnd(val) if val else "—"
            elif key == "female_family":
                val = line.get("price_female_family") or 0
                cell_text = fmt_vnd(val) if val else "—"
            else:
                val = (line.get("extra_prices_json") or {}).get(key) or 0
                cell_text = fmt_vnd(val) if val else "—"
            _write_cell_text(tr.cells[cell_idx], cell_text, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    # Totals rows
    from apps.contract.services.document_payloads import fmt_vnd
    totals = pkg_payload.get("totals") or {}
    by_col = totals.get("by_col") or {}

    # Row giá niêm yết tổng
    tr = table.add_row()
    apply_widths(tr)
    tr.height = Cm(0.85)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    merged_label = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_label, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(merged_label, "GIÁ TRỊ GÓI THEO NIÊM YẾT (VND/NGƯỜI)",
                     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9, color="D35400", shading="F3F3F3")
    for ci, col in enumerate(columns):
        key = col.get("key", "")
        val = by_col.get(key, {}).get("base_per_person", 0)
        _write_cell_text(tr.cells[3 + ci], fmt_vnd(val) if val else "—",
                         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, italic=True, size=9, color="D35400", shading="F3F3F3")

    # Row giá ưu đãi
    tr = table.add_row()
    apply_widths(tr)
    tr.height = Cm(0.85)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    merged_disc = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_disc, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(merged_disc, "GIÁ ƯU ĐÃI DÀNH CHO QUÝ CTY (VND/NGƯỜI)",
                     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9, color="FFFFFF", shading="FFC107")
    for ci, col in enumerate(columns):
        key = col.get("key", "")
        val = by_col.get(key, {}).get("per_person", 0)
        _write_cell_text(tr.cells[3 + ci], fmt_vnd(val) if val else "—",
                         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9, color="000000", shading="FFC107")

    # Row số lượng
    tr = table.add_row()
    apply_widths(tr)
    tr.height = Cm(0.75)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    merged_count = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_count, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(merged_count, "SỐ LƯỢNG (NGƯỜI)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9, shading="FAFAFA")
    for ci, col in enumerate(columns):
        key = col.get("key", "")
        count = by_col.get(key, {}).get("count", col.get("count", 0))
        _write_cell_text(tr.cells[3 + ci], str(count) if count else "—",
                         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9, shading="FAFAFA")

    # Chèn bảng vào document sau pkg_para
    pkg_para.addnext(table._tbl)
    return table._tbl


class _MiniQuillDocxParser(HTMLParser):
    BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "li"}

    def __init__(self):
        super().__init__()
        self.blocks = []
        self.current_block = None
        self.inline_stack = []
        self.list_stack = []

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs or [])
        tag = (tag or "").lower()

        if tag in ("ul", "ol"):
            self.list_stack.append({"type": tag, "index": 1})
            return

        if tag == "li":
            self._start_block(tag, attrs, force_new=True)
            return

        if tag in self.BLOCK_TAGS:
            self._start_block(tag, attrs, force_new=True)
            return

        if tag == "br":
            self._append_text("\n")
            return

        self.inline_stack.append((tag, self._inline_style_from(tag, attrs)))

    def handle_endtag(self, tag):
        tag = (tag or "").lower()

        if tag in ("ul", "ol"):
            if self.list_stack:
                self.list_stack.pop()
            return

        if tag in self.BLOCK_TAGS:
            self._flush_block()
            return

        for i in range(len(self.inline_stack) - 1, -1, -1):
            if self.inline_stack[i][0] == tag:
                self.inline_stack.pop(i)
                break

    def handle_data(self, data):
        if data is None:
            return
        text = unescape(data).replace("\xa0", " ")
        if not text:
            return
        self._append_text(text)

    def close(self):
        super().close()
        self._flush_block()

    def _parse_classes(self, attrs):
        class_list = (attrs.get("class") or "").split()
        align = None
        indent = 0
        size = None

        if "ql-align-center" in class_list:
            align = "center"
        elif "ql-align-right" in class_list:
            align = "right"
        elif "ql-align-justify" in class_list:
            align = "justify"

        for i in range(1, 9):
            if f"ql-indent-{i}" in class_list:
                indent = i
                break

        if "ql-size-small" in class_list:
            size = 9
        elif "ql-size-large" in class_list:
            size = 13
        elif "ql-size-huge" in class_list:
            size = 16

        return {
            "align": align,
            "indent": indent,
            "size": size,
        }

    def _block_style_from(self, tag, attrs):
        cls = self._parse_classes(attrs)
        style = {
            "tag": tag,
            "align": cls["align"],
            "indent": cls["indent"] or 0,
            "size": cls["size"] or 10.5,
            "bold": False,
            "italic": False,
            "list_prefix": "",
        }

        if tag == "h1":
            style.update({"size": 16, "bold": True})
        elif tag == "h2":
            style.update({"size": 14, "bold": True})
        elif tag == "h3":
            style.update({"size": 12.5, "bold": True})
        elif tag == "h4":
            style.update({"size": 11.5, "bold": True})
        elif tag in {"h5", "h6"}:
            style.update({"size": 11, "bold": True})
        elif tag == "blockquote":
            style.update({"italic": True})

        if tag == "li" and self.list_stack:
            current = self.list_stack[-1]
            if current["type"] == "ul":
                style["list_prefix"] = "• "
            else:
                style["list_prefix"] = f"{current['index']}. "
                current["index"] += 1

        return style

    def _inline_style_from(self, tag, attrs):
        attrs = attrs or {}
        cls = self._parse_classes(attrs)

        style = {
            "bold": False,
            "italic": False,
            "underline": False,
            "strike": False,
            "size": cls["size"],
            "link": None,
        }

        if tag in ("strong", "b"):
            style["bold"] = True
        if tag in ("em", "i"):
            style["italic"] = True
        if tag == "u":
            style["underline"] = True
        if tag in ("s", "strike"):
            style["strike"] = True
        if tag == "a":
            style["link"] = attrs.get("href")

        return style

    def _start_block(self, tag, attrs, force_new=False):
        if force_new and self.current_block and self.current_block["runs"]:
            self._flush_block()

        if self.current_block is None:
            st = self._block_style_from(tag, attrs)
            self.current_block = {
                "tag": st["tag"],
                "align": st["align"],
                "indent": st["indent"],
                "size": st["size"],
                "bold": st["bold"],
                "italic": st["italic"],
                "list_prefix": st["list_prefix"],
                "runs": [],
            }

    def _ensure_default_block(self):
        if self.current_block is None:
            self._start_block("p", {}, force_new=False)

    def _current_inline_style(self):
        merged = {
            "bold": False,
            "italic": False,
            "underline": False,
            "strike": False,
            "size": None,
            "link": None,
        }
        for _, item in self.inline_stack:
            merged["bold"] = merged["bold"] or item.get("bold", False)
            merged["italic"] = merged["italic"] or item.get("italic", False)
            merged["underline"] = merged["underline"] or item.get("underline", False)
            merged["strike"] = merged["strike"] or item.get("strike", False)
            merged["link"] = merged["link"] or item.get("link")
            if item.get("size"):
                merged["size"] = item["size"]
        return merged

    def _append_text(self, text):
        self._ensure_default_block()
        if not text:
            return
        inline = self._current_inline_style()
        self.current_block["runs"].append({
            "text": text,
            "bold": inline["bold"],
            "italic": inline["italic"],
            "underline": inline["underline"],
            "strike": inline["strike"],
            "size": inline["size"],
            "link": inline["link"],
        })

    def _flush_block(self):
        if not self.current_block:
            return

        has_real_text = any((r.get("text") or "").strip() for r in self.current_block["runs"])
        if has_real_text:
            self.blocks.append(self.current_block)

        self.current_block = None


def _apply_docx_run_style(run, *, bold=False, italic=False, underline=False, strike=False, size=10.5, color=None, font_name="Arial"):
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.strike = strike
    run.font.size = Pt(size)
    run.font.name = font_name

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)

    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _html_align_to_docx(value):
    if value == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if value == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if value == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def _html_to_docx_paragraphs(doc: Document, html: str, insert_after_element):
    """
    Chuyển HTML từ Quill -> paragraph/run trong docx, giữ được phần lớn định dạng:
    - heading
    - bold / italic / underline / strike
    - xuống dòng
    - ul / ol
    - align
    - indent
    - ql-size-small / large / huge
    """
    if not html or not html.strip():
        return insert_after_element

    parser = _MiniQuillDocxParser()
    parser.feed(html)
    parser.close()

    last_el = insert_after_element

    for block in parser.blocks:
        para = doc.add_paragraph()
        para.alignment = _html_align_to_docx(block.get("align"))
        para.paragraph_format.left_indent = Cm(0.6 * int(block.get("indent") or 0))

        tag = block.get("tag")
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            _set_paragraph_spacing(para, before=2, after=3, line=1.15)
        elif tag == "blockquote":
            _set_paragraph_spacing(para, before=1, after=2, line=1.15)
            para.paragraph_format.left_indent = Cm(0.8 + (0.6 * int(block.get("indent") or 0)))
        else:
            _set_paragraph_spacing(para, before=0, after=2, line=1.12)

        prefix = block.get("list_prefix") or ""
        if prefix:
            run = para.add_run(prefix)
            _apply_docx_run_style(
                run,
                bold=False,
                italic=False,
                underline=False,
                strike=False,
                size=block.get("size") or 10.5,
            )

        for item in block.get("runs", []):
            text = item.get("text") or ""
            if not text:
                continue

            parts = text.split("\n")
            for idx, part in enumerate(parts):
                if idx > 0:
                    para.add_run().add_break()
                if not part:
                    continue

                run = para.add_run(part)
                _apply_docx_run_style(
                    run,
                    bold=block.get("bold", False) or item.get("bold", False),
                    italic=block.get("italic", False) or item.get("italic", False),
                    underline=item.get("underline", False),
                    strike=item.get("strike", False),
                    size=item.get("size") or block.get("size") or 10.5,
                    color="0D6EFD" if item.get("link") else None,
                )

        last_el.addnext(para._element)
        last_el = para._element

    return last_el


def _inject_packages_at_marker(doc: Document, payload: dict) -> bool:
    marker = _find_table_marker(doc)
    if not marker:
        return False

    kind = marker[0]
    if kind != "document":
        return False  # Only handle document-level for multi-package

    paragraph = marker[1]
    paragraph.text = paragraph.text.replace(QUOTATION_TABLE_MARKER, "").strip()
    if not paragraph.text.strip():
        insert_point = paragraph._element
    else:
        insert_point = paragraph._element

    packages = payload.get("packages") or []
    last_el = insert_point

    for pkg in packages:
        last_tbl = _append_package_table(last_el, doc, pkg)
        last_el = last_tbl

    # Extra content (HTML từ Quill) sau tất cả bảng
    extra_content = (payload.get("quotation") or {}).get("extra_content") or ""
    if extra_content and extra_content.strip():
        last_el = _html_to_docx_paragraphs(doc, extra_content, last_el)

    if not paragraph.text.strip():
        _remove_paragraph(paragraph)
    return True


def _inject_table_at_marker(doc: Document, payload: dict) -> bool:
    """Legacy single-package injection."""
    marker = _find_table_marker(doc)
    if not marker:
        return False

    kind = marker[0]
    if kind == "document":
        paragraph = marker[1]
        paragraph.text = paragraph.text.replace(QUOTATION_TABLE_MARKER, "").strip()
        table = doc.add_table(rows=1, cols=6)
        table._tbl.getparent().remove(table._tbl)
        paragraph._p.addnext(table._tbl)
        proxy = type("TC", (), {"add_table": lambda self, rows, cols: table})()
        _append_quotation_table(proxy, payload)

        # Extra content
        extra_content = (payload.get("quotation") or {}).get("extra_content") or ""
        if extra_content:
            _html_to_docx_paragraphs(doc, extra_content, table._tbl)

        if not paragraph.text.strip():
            _remove_paragraph(paragraph)
        return True

    if kind == "cell":
        cell, paragraph = marker[1], marker[2]
        paragraph.text = paragraph.text.replace(QUOTATION_TABLE_MARKER, "").strip()
        _clear_cell(cell)
        _append_quotation_table(cell, payload)
        return True

    return False


def _append_quotation_table(container, payload: dict):
    """Legacy 6-column quotation table (backward compat)."""
    table = container.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, size=4, color="999999")
    _set_table_width(table, 100)
    _set_table_fixed_layout(table)

    widths_cm = [0.95, 7.25, 1.85, 2.55, 2.85, 2.85]
    _set_table_grid(table, widths_cm)

    def apply_widths(row):
        for idx, width_cm in enumerate(widths_cm):
            _set_cell_width(row.cells[idx], width_cm)

    for row in table.rows:
        apply_widths(row)

    header = table.rows[0]
    apply_widths(header)
    _set_row_repeat_as_header(header)
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    headers = ["STT", "KHÁM TỔNG QUÁT /\nGENERAL EXAMINATION", "GIÁ ƯU ĐÃI\n(VND)", "NAM", "NỮ\nĐỘC THÂN", "NỮ\nGIA ĐÌNH"]
    for idx, label in enumerate(headers):
        _write_cell_text(header.cells[idx], label, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9.5, color="FFFFFF", shading="2F5AA8")

    for row in payload.get("display_rows", []):
        if row["row_type"] in ("group", "subgroup"):
            tr = table.add_row()
            apply_widths(tr)
            tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            merged = tr.cells[0].merge(tr.cells[5])
            is_sub = row["row_type"] == "subgroup"
            _write_cell_text(merged, row["label"], align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, italic=is_sub,
                             size=10 if is_sub else 10.5, color=None if is_sub else "FFFFFF",
                             shading="D6EAF8" if is_sub else "2F5AA8")
            continue
        line = row.get("line") or {}
        tr = table.add_row()
        apply_widths(tr)
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        svc_text = line.get("item_name", "")
        if line.get("description"):
            svc_text = f"{svc_text}\n{line['description']}"
        _write_cell_text(tr.cells[0], line.get("stt", ""), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _write_cell_text(tr.cells[1], svc_text, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
        _write_cell_text(tr.cells[2], line.get("display_unit_price") or "", align=WD_ALIGN_PARAGRAPH.CENTER, bold=bool(line.get("display_unit_price")), size=10, shading="D9E5F6")
        _write_cell_text(tr.cells[3], line.get("male_mark", "–"), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _write_cell_text(tr.cells[4], line.get("female_single_mark", "–"), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _write_cell_text(tr.cells[5], line.get("female_family_mark", "–"), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    totals = payload.get("totals", {})
    from apps.contract.services.document_payloads import fmt_vnd

    tr = table.add_row()
    apply_widths(tr)
    tr.height = Cm(1.15)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    merged = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(merged, "GIÁ TRỊ GÓI KHÁM THEO GIÁ NIÊM YẾT TẠI PHÒNG KHÁM\n(VND/NGƯỜI)",
                     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9.5, color="FFFFFF", shading="2F5AA8")
    _write_cell_text(tr.cells[3], totals.get("base_total_male_display") or "0", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, color="FFFFFF", shading="2F5AA8")
    _write_cell_text(tr.cells[4], totals.get("base_total_female_single_display") or "0", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, color="FFFFFF", shading="2F5AA8")
    _write_cell_text(tr.cells[5], totals.get("base_total_female_family_display") or "0", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, color="FFFFFF", shading="2F5AA8")
    _set_cell_vertical_center(merged)

    tr = table.add_row()
    apply_widths(tr)
    tr.height = Cm(0.95)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    merged_disc = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_disc, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(merged_disc, "GIÁ ƯU ĐÃI DÀNH CHO QUÝ CTY (VND)",
                     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, shading="FFF200")
    _write_cell_text(tr.cells[3], totals.get("discounted_total_male_display") or "0", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, shading="FFF200")
    _write_cell_text(tr.cells[4], totals.get("discounted_total_female_single_display") or "0", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, shading="FFF200")
    _write_cell_text(tr.cells[5], totals.get("discounted_total_female_family_display") or "0", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, shading="FFF200")
    _set_cell_vertical_center(merged_disc)
    return table


def render_quotation_docx(*, payload: dict, output_path: str, template_path: str | None = None):
    if template_path:
        template_file = Path(template_path)
        if template_file.exists() and template_file.is_file():
            doc = Document(str(template_file))
            replacements = _build_replacements(payload)
            _replace_text_everywhere(doc, replacements)

            if payload.get("multi_package") and payload.get("packages"):
                inserted = _inject_packages_at_marker(doc, payload)
            else:
                inserted = _inject_table_at_marker(doc, payload)

            if not inserted:
                doc.add_paragraph("")
                _append_quotation_table(doc, payload)
            doc.save(output_path)
            return output_path

    return create_default_quotation_docx(payload, output_path)


def create_default_quotation_docx(payload: dict, output_path: str):
    quotation = payload.get("quotation", {})
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("BẢNG BÁO GIÁ KHÁM SỨC KHỎE")
    _style_runs(p, bold=True, size=20, color="2F5AA8")
    _set_paragraph_spacing(p, after=4)

    for line in [
        f"Kính gửi: {quotation.get('contact_name', '')}",
        f"Công ty: {quotation.get('company_name', '')}",
        f"Địa chỉ: {quotation.get('company_address', '')}",
    ]:
        p = doc.add_paragraph(line)
        _style_runs(p, size=11)
        _set_paragraph_spacing(p, after=2)

    if payload.get("multi_package") and payload.get("packages"):
        for pkg in payload["packages"]:
            p = doc.add_paragraph(pkg.get("name", "Gói khám"))
            _style_runs(p, bold=True, size=13, color="1A5276")
            _set_paragraph_spacing(p, before=8, after=4)
            pkg_proxy = type("C", (), {"add_table": lambda self, rows, cols, style=None: doc.add_table(rows, cols)})()
            # Reuse _append_package_table logic simplified
            for dr in pkg.get("display_rows", []):
                pass  # simplified: just add table via _append_package_table
    else:
        _append_quotation_table(doc, payload)

    extra = quotation.get("extra_content") or ""
    if extra:
        p = doc.add_paragraph("")
        _html_to_docx_paragraphs(doc, extra, p._element)

    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Contract renderer (unchanged from original)
# ─────────────────────────────────────────────────────────────────────────────

def _build_contract_replacements(payload: dict) -> dict:
    return {
        "{{ contract_number }}":       payload.get("contract_number", ""),
        "{{ contract_number_full }}":  payload.get("contract_number_full", ""),
        "{{ issue_day }}":             payload.get("issue_day", ""),
        "{{ issue_month }}":           payload.get("issue_month", ""),
        "{{ issue_year }}":            payload.get("issue_year", ""),
        "{{ issue_date }}":            payload.get("issue_date", ""),
        "{{ issue_date_vi }}":         payload.get("issue_date_vi", ""),
        "{{ company_name }}":          payload.get("company_name", ""),
        "{{ company_address }}":       payload.get("company_address", ""),
        "{{ company_phone }}":         payload.get("company_phone", ""),
        "{{ company_tax_code }}":      payload.get("company_tax_code", ""),
        "{{ signer_b_name }}":         payload.get("signer_b_name", ""),
        "{{ signer_b_title }}":        payload.get("signer_b_title", ""),
        "{{ signer_a_name }}":         payload.get("signer_a_name", ""),
        "{{ signer_a_title }}":        payload.get("signer_a_title", ""),
        "{{ period_text }}":           payload.get("period_text", ""),
        "{{ start_date }}":            payload.get("start_date", ""),
        "{{ end_date }}":              payload.get("end_date", ""),
        "{{ reception_from_date }}":   payload.get("reception_from_date", ""),
        "{{ blood_time_from }}":       payload.get("blood_time_from", ""),
        "{{ blood_time_to }}":         payload.get("blood_time_to", ""),
        "{{ blood_time_text }}":       payload.get("blood_time_text", ""),
        "{{ blood_collection_location }}": payload.get("blood_collection_location", ""),
        "{{ male_count }}":            str(payload.get("male_count", 0)),
        "{{ female_single_count }}":   str(payload.get("female_single_count", 0)),
        "{{ female_family_count }}":   str(payload.get("female_family_count", 0)),
        "{{ total_pax }}":             str(payload.get("total_pax", 0)),
        "{{ subtotal_male }}":         payload.get("subtotal_male", ""),
        "{{ subtotal_female_single }}": payload.get("subtotal_female_single", ""),
        "{{ subtotal_female_family }}": payload.get("subtotal_female_family", ""),
        "{{ grand_total }}":           payload.get("grand_total", ""),
        "{{ deposit_pct }}":           payload.get("deposit_pct", ""),
        "{{ deposit_amount }}":        payload.get("deposit_amount", ""),
        "{{ deposit_amount_words }}":  payload.get("deposit_amount_words", ""),
        "{{ deposit_deadline }}":      payload.get("deposit_deadline", ""),
        "{{ settlement_days }}":       payload.get("settlement_days", ""),
        "{{ contract_note }}":         payload.get("contract_note", ""),
        "{{ note }}":                  payload.get("note", ""),
    }


def _append_contract_catalog_section(insert_after_element, doc: Document, section: dict):
    """
    Render 1 section catalog hợp đồng từ schema chung `catalog_sections`.
    Trả về XML element cuối cùng vừa chèn để chain tiếp.
    """
    columns = section.get("columns") or [
        {"key": "male", "label": "NAM"},
        {"key": "female_single", "label": "NỮ ĐỘC THÂN"},
        {"key": "female_family", "label": "NỮ GIA ĐÌNH"},
    ]
    rows = section.get("rows") or []
    title = (section.get("title") or "").strip()

    num_cols = len(columns)
    total_cols = 3 + num_cols
    widths_cm = _compute_col_widths(num_cols)

    anchor_el = insert_after_element

    if title:
        pkg_para = OxmlElement("w:p")
        pkg_pPr = OxmlElement("w:pPr")
        pkg_spacing = OxmlElement("w:spacing")
        pkg_spacing.set(qn("w:before"), "120")
        pkg_spacing.set(qn("w:after"), "60")
        pkg_pPr.append(pkg_spacing)
        pkg_para.append(pkg_pPr)

        pkg_r = OxmlElement("w:r")
        pkg_rPr = OxmlElement("w:rPr")
        pkg_b = OxmlElement("w:b")
        pkg_sz = OxmlElement("w:sz")
        pkg_sz.set(qn("w:val"), "22")
        pkg_color = OxmlElement("w:color")
        pkg_color.set(qn("w:val"), "1A5276")
        pkg_rPr.append(pkg_b)
        pkg_rPr.append(pkg_sz)
        pkg_rPr.append(pkg_color)
        pkg_r.append(pkg_rPr)

        pkg_t = OxmlElement("w:t")
        pkg_t.text = title
        pkg_r.append(pkg_t)
        pkg_para.append(pkg_r)

        anchor_el.addnext(pkg_para)
        anchor_el = pkg_para

    table = doc.add_table(rows=1, cols=total_cols)
    table._tbl.getparent().remove(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, size=4, color="999999")
    _set_table_width(table, 100)
    _set_table_fixed_layout(table)
    _set_table_grid(table, widths_cm)

    def apply_widths(row):
        for idx, width_cm in enumerate(widths_cm):
            _set_cell_width(row.cells[idx], width_cm)

    for row in table.rows:
        apply_widths(row)

    header = table.rows[0]
    apply_widths(header)
    _set_row_repeat_as_header(header)
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    header_labels = ["STT", "KHÁM TỔNG QUÁT /\nGENERAL EXAMINATION", "GIÁ ƯU ĐÃI\n(VNĐ)"]
    for col in columns:
        header_labels.append((col.get("label") or "").upper())

    for idx, label in enumerate(header_labels):
        _write_cell_text(
            header.cells[idx],
            label,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            size=9.5,
            color="FFFFFF",
            shading="2F5AA8",
        )

    for row_data in rows:
        row_type = row_data.get("row_type")

        if row_type in ("group", "subgroup"):
            tr = table.add_row()
            apply_widths(tr)
            tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            merged = tr.cells[0].merge(tr.cells[total_cols - 1])
            is_sub = row_type == "subgroup"
            _write_cell_text(
                merged,
                row_data.get("label", ""),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
                italic=is_sub,
                size=10 if is_sub else 10.5,
                color=None if is_sub else "FFFFFF",
                shading="D6EAF8" if is_sub else "2F5AA8",
            )
            continue

        line = row_data.get("line") or {}
        tr = table.add_row()
        apply_widths(tr)
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        _write_cell_text(
            tr.cells[0],
            line.get("stt", ""),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=10,
        )

        svc_text = line.get("item_name", "")
        if line.get("description"):
            svc_text = f"{svc_text}\n{line['description']}"
        _write_cell_text(
            tr.cells[1],
            svc_text,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            size=10,
        )

        _write_cell_text(
            tr.cells[2],
            line.get("display_unit_price") or "–",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=bool(line.get("display_unit_price")),
            size=10,
            shading="D9E5F6",
        )

        for ci, col in enumerate(columns):
            key = col.get("key", "")
            cell_text = "–"

            if key == "male":
                cell_text = line.get("male_mark", "–")
            elif key == "female_single":
                cell_text = line.get("female_single_mark", "–")
            elif key == "female_family":
                cell_text = line.get("female_family_mark", "–")
            else:
                extra_marks = line.get("extra_column_marks") or {}
                extra_prices = line.get("extra_prices_json") or {}
                cell_text = extra_marks.get(key) or (fmt_vnd(extra_prices.get(key)) if extra_prices.get(key) else "–")

            _write_cell_text(
                tr.cells[3 + ci],
                cell_text or "–",
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size=10,
            )

    # Row số lượng
    tr = table.add_row()
    apply_widths(tr)
    tr.height = Cm(0.75)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    merged_count = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_count, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(
        merged_count,
        "SỐ LƯỢNG NHÂN VIÊN",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=9,
        shading="FAFAFA",
    )
    for ci, col in enumerate(columns):
        _write_cell_text(
            tr.cells[3 + ci],
            str(_to_int(col.get("count")) or 0),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            size=9,
            shading="FAFAFA",
        )

    # Row giá niêm yết
    tr = table.add_row()
    apply_widths(tr)
    tr.height = Cm(0.85)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    merged_label = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_label, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(
        merged_label,
        "GIÁ TRỊ GÓI KHÁM THEO GIÁ NIÊM YẾT (VND/NGƯỜI)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=9,
        color="D35400",
        shading="F3F3F3",
    )
    for ci, col in enumerate(columns):
        _write_cell_text(
            tr.cells[3 + ci],
            col.get("base_per_person_display") or "–",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            italic=True,
            size=9,
            color="D35400",
            shading="F3F3F3",
        )

    # Row giá ưu đãi
    tr = table.add_row()
    apply_widths(tr)
    tr.height = Cm(0.85)
    tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    merged_disc = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_disc, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(
        merged_disc,
        "GIÁ ƯU ĐÃI DÀNH CHO QUÝ CÔNG TY",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=9,
        color="FFFFFF",
        shading="2F7F8E",
    )
    for ci, col in enumerate(columns):
        _write_cell_text(
            tr.cells[3 + ci],
            col.get("per_person_display") or "–",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            size=9,
            color="FFFFFF",
            shading="2F7F8E",
        )

    anchor_el.addnext(table._tbl)
    return table._tbl


def _add_contract_catalog_table(container, payload: dict):
    table = container.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, size=4, color="999999")
    _set_table_width(table, 100)
    _set_table_fixed_layout(table)

    widths_cm = [0.95, 7.25, 1.85, 2.55, 2.85, 2.85]
    _set_table_grid(table, widths_cm)

    def apply_widths(row):
        for idx, width_cm in enumerate(widths_cm):
            _set_cell_width(row.cells[idx], width_cm)

    for row in table.rows:
        apply_widths(row)

    header = table.rows[0]
    apply_widths(header)
    _set_row_repeat_as_header(header)
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    headers = ["STT", "KHÁM TỔNG QUÁT /\nGENERAL EXAMINATION", "GIÁ ƯU ĐÃI\n(VND)", "NAM", "NỮ\nĐỘC THÂN", "NỮ\nGIA ĐÌNH"]
    for idx, label in enumerate(headers):
        _write_cell_text(header.cells[idx], label, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9.5, color="FFFFFF", shading="2F5AA8")

    display_rows = payload.get("display_rows") or []
    if not display_rows:
        for line in payload.get("lines", []):
            display_rows.append({"row_type": "item", "group_name": line.get("group_name") or "Khác", "subgroup_name": line.get("subgroup_name") or "", "line": line})

    current_group = None
    current_subgroup = None

    def add_group_row_local(label, *, shading="2F5AA8", italic=False):
        tr = table.add_row()
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        apply_widths(tr)
        merged = tr.cells[0].merge(tr.cells[5])
        _write_cell_text(merged, label, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, italic=italic,
                         size=10 if italic else 10.5, color=None if italic else "FFFFFF", shading=shading)

    def mark_text(enabled, price_type, note=None):
        if not enabled:
            return "–"
        if price_type == "free":
            return "Miễn phí"
        if price_type == "gift":
            return note or "TẶNG"
        return "✓"

    for row in display_rows:
        row_type = row.get("row_type")
        if row_type == "group":
            current_group = row.get("label") or "Khác"
            current_subgroup = None
            add_group_row_local(current_group)
            continue
        if row_type == "subgroup":
            current_subgroup = row.get("label") or ""
            if current_subgroup:
                add_group_row_local(current_subgroup, shading="D6EAF8", italic=True)
            continue
        line = row.get("line") or {}
        group_name = (line.get("group_name") or row.get("group_name") or "Khác").strip()
        subgroup_name = (line.get("subgroup_name") or row.get("subgroup_name") or "").strip()
        if group_name and group_name != current_group:
            current_group = group_name
            current_subgroup = None
            add_group_row_local(current_group)
        if subgroup_name and subgroup_name != current_subgroup:
            current_subgroup = subgroup_name
            add_group_row_local(current_subgroup, shading="D6EAF8", italic=True)
        tr = table.add_row()
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        apply_widths(tr)
        price_type = line.get("price_type") or "standard"
        note = line.get("note") or ""
        svc_text = line.get("item_name", "")
        if line.get("description"):
            svc_text = f"{svc_text}\n{line['description']}"
        _write_cell_text(tr.cells[0], line.get("stt", ""), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _write_cell_text(tr.cells[1], svc_text, align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
        _write_cell_text(tr.cells[2], line.get("display_unit_price") or "", align=WD_ALIGN_PARAGRAPH.CENTER, bold=bool(line.get("display_unit_price")), size=10, shading="D9E5F6" if price_type == "standard" else "F6C39A")
        _write_cell_text(tr.cells[3], line.get("male_mark") or mark_text(bool(line.get("for_male")), price_type, note), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _write_cell_text(tr.cells[4], line.get("female_single_mark") or mark_text(bool(line.get("for_female_single")), price_type, note), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _write_cell_text(tr.cells[5], line.get("female_family_mark") or mark_text(bool(line.get("for_female_family")), price_type, note), align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

    from apps.contract.services.document_payloads import fmt_vnd

    for label, key_base, key_disc in [
        ("GIÁ TRỊ GÓI KHÁM THEO GIÁ NIÊM YẾT (VND/NGƯỜI)", "base_total", "discounted_unit"),
        ("GIÁ ƯU ĐÃI DÀNH CHO QUÝ CÔNG TY", "subtotal", "subtotal"),
    ]:
        tr = table.add_row()
        apply_widths(tr)
        tr.height = Cm(0.95)
        tr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        merged = tr.cells[0].merge(tr.cells[2])
        _set_cell_width(merged, widths_cm[0] + widths_cm[1] + widths_cm[2])
        is_disc = "ƯU ĐÃI" in label
        shd = "2F7F8E" if is_disc else "F3F3F3"
        clr = "FFFFFF" if is_disc else "D35400"
        _write_cell_text(merged, label, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9.5, color=clr, shading=shd)
        if is_disc:
            vals = [payload.get("subtotal_male", "–"), payload.get("subtotal_female_single", "–"), payload.get("subtotal_female_family", "–")]
        else:
            vals = [payload.get("base_total_male_display", "–"), payload.get("base_total_female_single_display", "–"), payload.get("base_total_female_family_display", "–")]
        for ci, val in enumerate(vals):
            _write_cell_text(tr.cells[3 + ci], val or "–", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, color=clr if not is_disc else "FFFFFF", shading=shd)

    # Row grand total
    tr = table.add_row()
    apply_widths(tr)
    merged_l = tr.cells[0].merge(tr.cells[2])
    _set_cell_width(merged_l, widths_cm[0] + widths_cm[1] + widths_cm[2])
    _write_cell_text(merged_l, "GIÁ TRỊ HỢP ĐỒNG (VND)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10.5, shading="F3F3F3")
    merged_v = tr.cells[3].merge(tr.cells[5])
    _set_cell_width(merged_v, widths_cm[3] + widths_cm[4] + widths_cm[5])
    _write_cell_text(merged_v, payload.get("grand_total") or "0", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11, shading="F3F3F3")
    _set_cell_vertical_center(merged_l)
    _set_cell_vertical_center(merged_v)

    return table


def _find_contract_catalog_marker(doc: Document):
    for paragraph in doc.paragraphs:
        if CONTRACT_CATALOG_TABLE_MARKER in paragraph.text:
            return ("document", paragraph)
    for table in _iter_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if CONTRACT_CATALOG_TABLE_MARKER in paragraph.text:
                        return ("cell", cell, paragraph)
    return None


def _inject_contract_catalog_sections(doc: Document, payload: dict) -> bool:
    marker = _find_contract_catalog_marker(doc)
    if not marker:
        return False

    sections = payload.get("catalog_sections") or []
    if not sections:
        return False

    kind = marker[0]
    if kind != "document":
        return False

    paragraph = marker[1]
    paragraph.text = paragraph.text.replace(CONTRACT_CATALOG_TABLE_MARKER, "").strip()

    last_el = paragraph._element
    for section in sections:
        last_el = _append_contract_catalog_section(last_el, doc, section)

    extra = payload.get("extra_content") or ""
    if extra:
        _html_to_docx_paragraphs(doc, extra, last_el)

    if not paragraph.text.strip():
        _remove_paragraph(paragraph)
    return True


def _inject_contract_catalog(doc: Document, payload: dict) -> bool:
    """
    Backward-compat wrapper.
    Ưu tiên schema mới `catalog_sections`, fallback về bảng legacy cũ nếu cần.
    """
    if payload.get("catalog_sections"):
        return _inject_contract_catalog_sections(doc, payload)

    marker = _find_contract_catalog_marker(doc)
    if not marker:
        return False
    kind = marker[0]
    if kind == "document":
        paragraph = marker[1]
        paragraph.text = paragraph.text.replace(CONTRACT_CATALOG_TABLE_MARKER, "").strip()
        table = doc.add_table(rows=1, cols=6)
        table._tbl.getparent().remove(table._tbl)
        paragraph._p.addnext(table._tbl)
        proxy = type("TC", (), {"add_table": lambda self, rows, cols: table})()
        _add_contract_catalog_table(proxy, payload)

        extra = payload.get("extra_content") or ""
        if extra:
            _html_to_docx_paragraphs(doc, extra, table._tbl)

        if not paragraph.text.strip():
            _remove_paragraph(paragraph)
        return True

    if kind == "cell":
        cell, paragraph = marker[1], marker[2]
        paragraph.text = paragraph.text.replace(CONTRACT_CATALOG_TABLE_MARKER, "").strip()
        _clear_cell(cell)
        _add_contract_catalog_table(cell, payload)
        return True

    return False


def create_default_contract_docx(payload: dict, output_path: str) -> str:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HỢP ĐỒNG KHÁM SỨC KHỎE ĐỊNH KỲ")
    _style_runs(p, bold=True, size=16, color="1A5276")
    _set_paragraph_spacing(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Số: {payload.get('contract_number_full', '')}")
    _style_runs(p, size=12)
    _set_paragraph_spacing(p, after=12)

    if payload.get("catalog_sections"):
        last_el = doc.paragraphs[-1]._element
        for section_data in payload.get("catalog_sections", []):
            last_el = _append_contract_catalog_section(last_el, doc, section_data)
    elif payload.get("lines"):
        _add_contract_catalog_table(doc, payload)

    doc.add_paragraph()

    for line_txt in [
        f"Tổng giá trị hợp đồng: {payload.get('grand_total', '')} VNĐ",
        f"Đặt cọc {payload.get('deposit_pct', '')}: {payload.get('deposit_amount', '')} VNĐ",
    ]:
        p = doc.add_paragraph(line_txt)
        _style_runs(p, size=11)
        _set_paragraph_spacing(p, after=3)

    if payload.get("extra_content"):
        doc.add_paragraph("")
        _html_to_docx_paragraphs(doc, payload["extra_content"], doc.paragraphs[-1]._element)

    doc.save(output_path)
    return output_path


def render_contract_docx(*, payload: dict, output_path: str, template_path: str | None = None) -> str:
    if template_path:
        template_file = Path(template_path)
        if template_file.exists() and template_file.is_file():
            doc = Document(str(template_file))
            replacements = _build_contract_replacements(payload)
            _replace_text_everywhere(doc, replacements)

            if payload.get("catalog_sections"):
                _inject_contract_catalog_sections(doc, payload)
            elif payload.get("lines"):
                _inject_contract_catalog(doc, payload)

            doc.save(output_path)
            return output_path
    return create_default_contract_docx(payload, output_path)
